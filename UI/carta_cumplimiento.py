import sys, os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import re
import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
from config import RUTAS
from database import leer_tickets, registrar_documento_generado

# ==========================================================
# MESES
# ==========================================================
MESES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
}

def fecha_extensa(fecha):
    if pd.isna(fecha):
        return ""
    f = pd.to_datetime(fecha)
    return f"{f.day} de {MESES[f.month]} de {f.year}"

def limpiar_texto(valor):
    if pd.isna(valor):
        return ""
    texto = str(valor).strip()
    return re.sub(r' +', ' ', texto)

# ==========================================================
# REEMPLAZO SEGURO EN WORD
# ==========================================================
def _reemplazar_en_parrafo(parrafo, campos, marcador_fmt):
    texto_original = "".join(run.text for run in parrafo.runs)
    texto_nuevo = texto_original
    for k, v in campos.items():
        marcador = marcador_fmt(k)
        texto_nuevo = texto_nuevo.replace(marcador, v)
    if texto_original != texto_nuevo and parrafo.runs:
        parrafo.runs[0].text = texto_nuevo
        for run in parrafo.runs[1:]:
            run.text = ""

def _reemplazar_en_tabla(tabla, campos, marcador_fmt):
    for fila in tabla.rows:
        for celda in fila.cells:
            for parrafo in celda.paragraphs:
                _reemplazar_en_parrafo(parrafo, campos, marcador_fmt)
            for tabla_anidada in celda.tables:
                _reemplazar_en_tabla(tabla_anidada, campos, marcador_fmt)

# ==========================================================
# GENERAR CARTA
# ==========================================================
def generar_carta(row, usuario: str = ""):
    doc = Document(RUTAS["plantilla_carta"])

    campos = {
        "N° CARTA DE CUMPLIMIENTO":                 limpiar_texto(row.get("N° CARTA DE CUMPLIMIENTO")),
        "USUARIO":                                  limpiar_texto(row.get("USUARIO")),
        "CORREO":                                   limpiar_texto(row.get("CORREO")),
        "RESOLUCION":                               limpiar_texto(row.get("RESOLUCION")),
        "TICKET":                                   limpiar_texto(row.get("TICKET")),
        "FECHA PROGRAMADA MANTENIMIENTO RES":       fecha_extensa(row.get("FECHA PROGRAMADA MANTENIMIENTO RES")),
        "FECHA DE CUMPLIMIENTO":                    fecha_extensa(row.get("FECHA DE CUMPLIMIENTO")),
        "FECHA DE ELABORACION CARTA DE CUMPLIMIENTO": fecha_extensa(row.get("FECHA DE ELABORACION CARTA DE CUMPLIMIENTO")),
    }

    marcador_fmt = lambda k: "{" + k + "}"

    for parrafo in doc.paragraphs:
        _reemplazar_en_parrafo(parrafo, campos, marcador_fmt)
    for tabla in doc.tables:
        _reemplazar_en_tabla(tabla, campos, marcador_fmt)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Registrar en historial
    if usuario:
        try:
            registrar_documento_generado(
                str(row.get("TICKET", "")),
                str(row.get("CLAVE_UNICA", "")),
                "CARTA GENERADA",
                usuario
            )
        except Exception:
            pass

    return buffer

# ==========================================================
# VISTA CARTA DE CUMPLIMIENTO
# ==========================================================
def mostrar_carta_cumplimiento():

    st.subheader("CARTA DE CUMPLIMIENTO")

    usuario = st.session_state.get("usuario_nombre", "")

    # Usar datos del session_state si están disponibles
    if "df" not in st.session_state or st.session_state.df.empty:
        st.warning("No hay consolidado disponible.")
        return

    df = st.session_state.df.copy()

    # Filtro obligatorio — solo tickets con fecha de resolución
    df = df[df["FECHA DE RESOLUCION"].notna()]

    if df.empty:
        st.info("No existen tickets que ameriten Carta de Cumplimiento.")
        return

    # Filtros
    with st.expander("🔍 Filtros de búsqueda", expanded=True):
        c1, c2, c3, c4, c5 = st.columns(5)

        ticket_f   = c1.text_input("TICKET", key="cc_ticket")
        solicitud_f = c2.multiselect("SOLICITUD",
            sorted(df["SOLICITUD"].astype(str).str.strip().str.upper().unique()),
            key="cc_solicitud")
        estado_f   = c3.multiselect("ESTADO",
            sorted(df["ESTADO"].astype(str).str.strip().str.upper().unique()),
            key="cc_estado")
        paso_f     = c4.multiselect("PASO CALIDAD",
            sorted(df["PASO A CALIDAD"].astype(str).str.strip().str.upper().unique()),
            key="cc_paso")
        fecha_f    = c5.date_input("FECHA DE AVERÍA", value=None, key="cc_fecha")

    if ticket_f:
        df = df[df["TICKET"].astype(str).str.contains(ticket_f, na=False)]
    if solicitud_f:
        df = df[df["SOLICITUD"].astype(str).str.strip().str.upper().isin(solicitud_f)]
    if estado_f:
        df = df[df["ESTADO"].astype(str).str.strip().str.upper().isin(estado_f)]
    if paso_f:
        df = df[df["PASO A CALIDAD"].astype(str).str.strip().str.upper().isin(paso_f)]
    if fecha_f:
        df = df[df["FECHA Y HORA DE LA AVERIA"].dt.date == fecha_f]

    if df.empty:
        st.info("No hay resultados con los filtros aplicados.")
        return

    st.divider()

    # Selección
    total = len(df)
    seleccionados = {}

    col_sel, col_info, col_btn = st.columns([1, 3, 2])
    with col_sel:
        seleccionar_todos = st.checkbox("Seleccionar todos", key="cc_sel_todos")
    with col_info:
        st.caption(f"{total} ticket(s) disponibles para carta de cumplimiento")
    with col_btn:
        generar_clicked = st.button(
            "📄 Generar cartas de cumplimiento",
            key="cc_btn_generar",
            use_container_width=True
        )

    if generar_clicked:
        st.session_state["cc_documentos_generados"] = True

    st.divider()
    # Encabezados de columnas
    _, enc = st.columns([0.3, 9.7])
    with enc:
        e1, e2, e3, e4, e5, e6 = st.columns([1, 1.5, 1.5, 1.5, 1.5, 1.5])
        e1.markdown("**TICKET**")
        e2.markdown("**FECHA AVERÍA**")
        e3.markdown("**SOLICITUD**")
        e4.markdown("**ESTADO**")
        e5.markdown("**FECHA RESOLUCIÓN**")
        e6.markdown("**FECHA RESTABLECIMIENTO**")
    st.divider()

    # Tabla con checkboxes
    for idx, row in df.reset_index(drop=True).iterrows():
        with st.container():
            col_chk, col_datos = st.columns([0.3, 9.7])
            with col_chk:
                sel = st.checkbox("", value=seleccionar_todos, key=f"cc_chk_{row['TICKET']}_{idx}")
                seleccionados[idx] = (sel, row)
            with col_datos:
                c1, c2, c3, c4, c5, c6 = st.columns([1, 1.5, 1.5, 1.5, 1.5, 1.5])
                c1.write(f"**{row['TICKET']}**")
                c2.write(pd.to_datetime(row["FECHA Y HORA DE LA AVERIA"]).strftime("%d/%m/%Y") if pd.notna(row["FECHA Y HORA DE LA AVERIA"]) else "")
                c3.write(row["SOLICITUD"])
                c4.write(row["ESTADO"])
                c5.write(pd.to_datetime(row.get("FECHA DE RESOLUCION")).strftime("%d/%m/%Y") if pd.notna(row.get("FECHA DE RESOLUCION")) else "")
                c6.write(pd.to_datetime(row.get("FECHA Y HORA DE RESTABLECIMIENTO DEL SERVICIO")).strftime("%d/%m/%Y") if pd.notna(row.get("FECHA Y HORA DE RESTABLECIMIENTO DEL SERVICIO")) else "")
        st.divider()

    # Botones de descarga
    tickets_sel = [(idx, row) for idx, (sel, row) in seleccionados.items() if sel]
    n_sel = len(tickets_sel)

    if st.session_state.get("cc_documentos_generados"):
        if n_sel == 0:
            st.warning("⚠️ No hay tickets seleccionados.")
        else:
            st.success(f"✅ {n_sel} carta(s) lista(s) para descargar:")
            for idx, row in tickets_sel:
                buffer = generar_carta(row, usuario)
                st.download_button(
                    label=f"⬇️ CARTA CUMPLIMIENTO - TICKET {row['TICKET']}",
                    data=buffer,
                    file_name=f"CARTA-CUMPLIMIENTO-TICKET-{row['TICKET']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"cc_dl_{row['TICKET']}_{idx}"
                )