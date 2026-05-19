import sys, os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import streamlit as st
from config import RUTAS
import pandas as pd
import os
from docx import Document
from io import BytesIO

# ==========================================================
# FECHA EXTENSA
# ==========================================================
MESES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
}

def fecha_extensa(fecha):
    if pd.isna(fecha):
        return ""
    f = pd.to_datetime(fecha)
    return f"{f.day} de {MESES[f.month]} de {f.year}"

# ==========================================================
# COLUMNAS A MOSTRAR EN LA TABLA
# ==========================================================
COLUMNAS_VISTA = [
    "TICKET",
    "FECHA Y HORA DE LA AVERIA",
    "SOLICITUD",
    "ESTADO",
    "PASO A CALIDAD",
    "FECHA PROGRAMADA MANTENIMIENTO RESOLUCIÓN",
    "FECHA PROGRAMADA MANTENIMIENTO RES",
    "N° OT",
    "FECHA DE OT",
    "SUPERVISOR MANTENIMIENTO",
    "COLABORADOR ENCARGADO",
    "DECLARACION DE RECLAMO",
    "QUIEN INDICO LA FECHA DE ATENCION EN LA RESOLUCION",
    "CALIDAD CERRADO?",
    "FECHA DE RESOLUCION",
    "FECHA DE NOTIFICACION DE LA RESOLUCION",
    "MEDIO DE NOTIFICACION DEL RECLAMO",
    "MOTIVO POR EL CUAL NO CUMPLIO CON EL PLAZO",
    "DESCRIPCION DE LA AVERIA",
    "FECHA DE CUMPLIMIENTO",
    "APLICA CARTA DE CUMPLIMIENTO",
    "N° CARTA DE CUMPLIMIENTO",
    "FECHA DE ELABORACION CARTA DE CUMPLIMIENTO",
    "FECHA REMISION CARTA DE CUMPLIMIENTO",
    "MEDIO DE NOTIFICACION LA CARTA DE CUMPLIMIENTO"
]

# ==========================================================
# CARGAR ÚLTIMO CONSOLIDADO
# ==========================================================
def cargar_ultimo_consolidado():
    carpeta = RUTAS["consolidado"]
    archivos = sorted(
        [f for f in os.listdir(carpeta) if f.startswith("CONSOLIDADO_")],
        reverse=True
    )

    if not archivos:
        return pd.DataFrame()

    return pd.read_excel(os.path.join(carpeta, archivos[0]))

# ==========================================================
# REEMPLAZO SEGURO — PRESERVA FORMATO DEL WORD
# ==========================================================
import re

# ==========================================================
# TRIM INTELIGENTE
# Elimina espacios al inicio y al final, y reduce espacios
# internos múltiples a uno solo, preservando nombres compuestos
# Ejemplo: "  LA   LIBERTAD  " → "LA LIBERTAD"
# ==========================================================
def _limpiar_campo(valor):
    if pd.isna(valor):
        return ""
    texto = str(valor).strip()
    texto = re.sub(r' +', ' ', texto)  # múltiples espacios → uno solo
    return texto


def _reemplazar_en_parrafo(parrafo, campos, marcador_fmt):
    """
    Reemplaza marcadores en un párrafo a nivel de run
    para no destruir el formato (negrita, fuente, tamaño, etc.)

    Si el marcador está partido entre varios runs, primero
    fusiona el texto del párrafo, reemplaza, y redistribuye
    el texto en los runs preservando el formato del primero.
    """
    for k, v in campos.items():
        marcador = marcador_fmt(k)

        # Verificar si el marcador existe en el párrafo completo
        texto_completo = "".join(run.text for run in parrafo.runs)
        if marcador not in texto_completo:
            continue

        # Intentar reemplazo directo run por run (caso simple)
        reemplazado = False
        for run in parrafo.runs:
            if marcador in run.text:
                run.text = run.text.replace(marcador, v)
                reemplazado = True

        # Si el marcador estaba partido entre runs (caso complejo)
        if not reemplazado:
            texto_nuevo = texto_completo.replace(marcador, v)
            # Limpiar todos los runs excepto el primero
            for i, run in enumerate(parrafo.runs):
                if i == 0:
                    run.text = texto_nuevo
                else:
                    run.text = ""


def _reemplazar_en_tabla(tabla, campos, marcador_fmt):
    """Recorre todas las celdas de una tabla y reemplaza marcadores."""
    for fila in tabla.rows:
        for celda in fila.cells:
            for parrafo in celda.paragraphs:
                _reemplazar_en_parrafo(parrafo, campos, marcador_fmt)
            # Tablas anidadas
            for tabla_anidada in celda.tables:
                _reemplazar_en_tabla(tabla_anidada, campos, marcador_fmt)


# ==========================================================
# GENERAR RESOLUCIÓN — PRESERVA FORMATO
# ==========================================================
def generar_resolucion(row):

    doc = Document(RUTAS["plantilla_res"])

    campos = {
        "TICKET":                              _limpiar_campo(row.get("TICKET")),
        "USUARIO":                             _limpiar_campo(row.get("USUARIO")),
        "LOCALIDAD":                           _limpiar_campo(row.get("LOCALIDAD")),
        "DISTRITO":                            _limpiar_campo(row.get("DISTRITO")),
        "PROVINCIA":                           _limpiar_campo(row.get("PROVINCIA")),
        "DEPARTAMENTO":                        _limpiar_campo(row.get("DEPARTAMENTO")),
        "FECHA LIMITE RESOLUCION":             fecha_extensa(row.get("FECHA LIMITE RESOLUCION")),
        "FECHA QUE INICIA RECLAMO DE CALIDAD": fecha_extensa(row.get("FECHA QUE INICIA RECLAMO DE CALIDAD")),
        "SEÑOR(A)":                            _limpiar_campo(row.get("SEÑOR(A)")),
        "IDENTIFICADO":                        _limpiar_campo(row.get("IDENTIFICADO")),
        "N° DOCUMENTO DNI / CE":               _limpiar_campo(row.get("N° DOCUMENTO DNI / CE")),
        "USUARIO(A)":                          _limpiar_campo(row.get("USUARIO(A)")),
        "CÓDIGO IIB":                          _limpiar_campo(row.get("CÓDIGO IIB")),
        "TIPO ENTIDAD":                        _limpiar_campo(row.get("TIPO ENTIDAD")),
        "CORREO":                              _limpiar_campo(row.get("CORREO"))
    }

    # Marcadores tipo «CAMPO»
    marcador_fmt = lambda k: f"«{k}»"

    # Reemplazar en párrafos del cuerpo
    for parrafo in doc.paragraphs:
        _reemplazar_en_parrafo(parrafo, campos, marcador_fmt)

    # Reemplazar en tablas
    for tabla in doc.tables:
        _reemplazar_en_tabla(tabla, campos, marcador_fmt)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================================
# PESTAÑA RESOLUCIÓN
# ==========================================================
def mostrar_resoluciones():

    st.subheader("RESOLUCIÓN")

    df = cargar_ultimo_consolidado()

    if df.empty:
        st.warning("No existe ningún consolidado disponible.")
        return

    # ======================================================
    # NORMALIZACIÓN (SOLO PARA FILTROS)
    # ======================================================
    df_f = df.copy()

    df_f["SOLICITUD"] = df_f["SOLICITUD"].astype(str).str.strip().str.upper()
    df_f["ESTADO"] = df_f["ESTADO"].astype(str).str.strip().str.upper()
    df_f["PASO A CALIDAD"] = df_f["PASO A CALIDAD"].astype(str).str.strip().str.upper()

    # ======================================================
    # FILTRO OBLIGATORIO (amerita resolución)
    # ======================================================
    df_f = df_f[
        (df_f["SOLICITUD"].isin([
            "SOLUCION REMOTA",
            "REPORTE POR PROBLEMAS DE CALIDAD Y AVERIA"
        ])) &
        (df_f["ESTADO"].isin(["CERRADO", "PENDIENTE"])) &
        (df_f["PASO A CALIDAD"] == "CALIDAD")
    ]

    if df_f.empty:
        st.info("No existen tickets que ameriten resolución.")
        return

    # ======================================================
    # FILTROS DEL USUARIO
    # ======================================================
    with st.expander("🔍 Filtros de búsqueda", expanded=True):

        c1, c2, c3, c4, c5 = st.columns(5)

        ticket_f = c1.text_input("TICKET", key="res_f_ticket")

        solicitud_f = c2.multiselect(
            "SOLICITUD",
            sorted(df_f["SOLICITUD"].unique()),
            key="res_f_solicitud"
        )

        estado_f = c3.multiselect(
            "ESTADO",
            sorted(df_f["ESTADO"].unique()),
            key="res_f_estado"
        )

        paso_f = c4.multiselect(
            "PASO CALIDAD",
            sorted(df_f["PASO A CALIDAD"].unique()),
            key="res_f_paso"
        )

        fecha_f = c5.date_input(
            "FECHA DE AVERÍA",
            value=None,
            key="res_f_fecha"
        )

    if ticket_f:
        df_f = df_f[df_f["TICKET"].astype(str).str.contains(ticket_f, na=False)]

    if solicitud_f:
        df_f = df_f[df_f["SOLICITUD"].isin(solicitud_f)]

    if estado_f:
        df_f = df_f[df_f["ESTADO"].isin(estado_f)]

    if paso_f:
        df_f = df_f[df_f["PASO A CALIDAD"].isin(paso_f)]

    if fecha_f:
        df_f = df_f[df_f["FECHA Y HORA DE LA AVERIA"].dt.date == fecha_f]

    if df_f.empty:
        st.info("No hay resultados con los filtros aplicados.")
        return

    st.divider()

    # ======================================================
    # SELECCIÓN DE TICKETS
    # ======================================================
    total = len(df_f)
    seleccionados = {}

    col_sel, col_info, col_btn = st.columns([1, 3, 2])
    with col_sel:
        seleccionar_todos = st.checkbox(
            "Seleccionar todos",
            key="res_sel_todos"
        )
    with col_info:
        st.caption(f"{total} ticket(s) disponibles para resolución")
    with col_btn:
        generar_clicked = st.button(
            "📄 Generar resoluciones",
            key="res_btn_generar",
            use_container_width=True
        )

    if generar_clicked:
        st.session_state["res_documentos_generados"] = True

    st.divider()

    # ======================================================
    # TABLA CON CHECKBOXES
    # ======================================================
    for idx, row_f in df_f.reset_index().iterrows():

        row = df.loc[row_f["index"]]

        with st.container():
            col_chk, col_datos = st.columns([0.3, 9.7])

            with col_chk:
                seleccionado = st.checkbox(
                    "",
                    value=seleccionar_todos,
                    key=f"res_chk_{row['TICKET']}_{idx}"
                )
                seleccionados[idx] = (seleccionado, row)

            with col_datos:
                c1, c2, c3, c4, c5 = st.columns([1, 1.5, 1.5, 1.5, 1.5])
                c1.write(f"**{row['TICKET']}**")
                c2.write(row["FECHA Y HORA DE LA AVERIA"])
                c3.write(row["SOLICITUD"])
                c4.write(row["ESTADO"])
                c5.write(row["PASO A CALIDAD"])

        st.divider()

    # ======================================================
    # BOTONES DE DESCARGA — aparecen al generar
    # ======================================================
    tickets_sel = [(idx, row) for idx, (sel, row) in seleccionados.items() if sel]
    n_sel = len(tickets_sel)

    if st.session_state.get("res_documentos_generados"):
        if n_sel == 0:
            st.warning("⚠️ No hay tickets seleccionados.")
        else:
            st.success(f"✅ {n_sel} resolución(es) lista(s) para descargar:")
            for idx, row in tickets_sel:
                buffer = generar_resolucion(row)
                st.download_button(
                    label=f"⬇️ RESOLUCION - TICKET {row['TICKET']}",
                    data=buffer,
                    file_name=f"RESOLUCION-TICKET-{row['TICKET']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"res_dl_{row['TICKET']}_{idx}"
                )